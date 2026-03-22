from __future__ import annotations

import base64
import io
import os
from pathlib import Path
from typing import BinaryIO

from pydantic import BaseModel

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".heic", ".heif"}
_CLAUDE_MEDIA_TYPES: dict[str, str] = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
}


class SheetData(BaseModel):
    name: str
    headers: list[str]
    rows: list[list[str]]


class ParsedFile(BaseModel):
    filename: str
    extension: str
    text: str
    sheets: list[SheetData] | None = None


class FileParseError(RuntimeError):
    pass


def parse_file(data: bytes, filename: str) -> ParsedFile:
    ext = Path(filename).suffix.lower()
    if ext == ".txt":
        return _parse_text(data, filename, ext)
    if ext == ".pdf":
        return _parse_pdf(data, filename, ext)
    if ext in (".xlsx", ".xls"):
        return _parse_excel(data, filename, ext)
    if ext in (".docx", ".doc"):
        return _parse_word(data, filename, ext)
    if ext in _IMAGE_EXTS:
        return _parse_image(data, filename, ext)
    raise FileParseError(f"Unsupported file type: {ext!r}")


def _parse_text(data: bytes, filename: str, ext: str) -> ParsedFile:
    text = data.decode("utf-8", errors="replace")
    return ParsedFile(filename=filename, extension=ext, text=text)


def _parse_pdf(data: bytes, filename: str, ext: str) -> ParsedFile:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise FileParseError("pypdf is required to parse PDF files") from exc

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[Page {i}]\n{page_text}")
    text = "\n\n".join(pages)
    return ParsedFile(filename=filename, extension=ext, text=text)


def _parse_excel(data: bytes, filename: str, ext: str) -> ParsedFile:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise FileParseError("openpyxl is required to parse Excel files") from exc

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheets: list[SheetData] = []
    text_parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_data = list(ws.iter_rows(values_only=True))
        if not rows_data:
            continue

        # First non-empty row is treated as headers
        headers = [str(c) if c is not None else "" for c in rows_data[0]]
        rows = [
            [str(c) if c is not None else "" for c in row]
            for row in rows_data[1:]
        ]
        sheets.append(SheetData(name=sheet_name, headers=headers, rows=rows))

        # Build human-readable text for LLM consumption
        sheet_lines = [f"[Sheet: {sheet_name}]", "\t".join(headers)]
        for row in rows:
            sheet_lines.append("\t".join(row))
        text_parts.append("\n".join(sheet_lines))

    text = "\n\n".join(text_parts)
    return ParsedFile(filename=filename, extension=ext, text=text, sheets=sheets)


def _parse_word(data: bytes, filename: str, ext: str) -> ParsedFile:
    try:
        from docx import Document
    except ImportError as exc:
        raise FileParseError("python-docx is required to parse Word files") from exc

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append("\t".join(cell.text.strip() for cell in row.cells))
        parts.append("\n".join(rows))

    text = "\n\n".join(parts)
    return ParsedFile(filename=filename, extension=ext, text=text)


# ---------------------------------------------------------------------------
# Images
# ---------------------------------------------------------------------------

def _heic_to_jpeg(data: bytes) -> bytes:
    """Convert HEIC/HEIF bytes to JPEG bytes using pillow-heif."""
    try:
        import pillow_heif
        from PIL import Image
    except ImportError as exc:
        raise FileParseError("pillow and pillow-heif are required to parse HEIC images") from exc

    pillow_heif.register_heif_opener()
    img = Image.open(io.BytesIO(data))
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="JPEG")
    return buf.getvalue()


def _parse_image(data: bytes, filename: str, ext: str) -> ParsedFile:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise FileParseError(
            "ANTHROPIC_API_KEY is required to parse image files"
        )

    try:
        import anthropic
    except ImportError as exc:
        raise FileParseError("anthropic is required to parse image files") from exc

    # HEIC/HEIF must be converted to JPEG first — Claude API doesn't accept them natively
    if ext in (".heic", ".heif"):
        image_data = _heic_to_jpeg(data)
        media_type = "image/jpeg"
    else:
        image_data = data
        media_type = _CLAUDE_MEDIA_TYPES.get(ext, "image/jpeg")

    encoded = base64.standard_b64encode(image_data).decode("utf-8")

    client = anthropic.Anthropic(api_key=api_key)
    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=2048,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": encoded,
                        },
                    },
                    {
                        "type": "text",
                        "text": (
                            "Extract all text visible in this image verbatim. "
                            "If there is no text, describe the image content concisely. "
                            "Do not add commentary or formatting beyond what is present."
                        ),
                    },
                ],
            }
        ],
    )

    text = message.content[0].text if message.content else ""
    return ParsedFile(filename=filename, extension=ext, text=text)
